"""
Unified text extraction from various file formats.
Supports: .pdf, .docx, .txt, .eml, .msg, .csv, .xlsx
"""

import csv
import io
import os
from pathlib import Path


def extract_text(file_path: str) -> str:
    """Extract plain text from a file based on its extension."""
    path = Path(file_path)
    ext = path.suffix.lower()

    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".doc": _extract_docx,
        ".txt": _extract_txt,
        ".md": _extract_txt,
        ".eml": _extract_eml,
        ".csv": _extract_csv,
        ".xlsx": _extract_xlsx_as_text,
        ".xls": _extract_xlsx_as_text,
    }

    extractor = extractors.get(ext)
    if extractor is None:
        # Fallback: try reading as plain text
        return _extract_txt(file_path)

    return extractor(file_path)


def _extract_pdf(file_path: str) -> str:
    try:
        import pdfplumber

        texts = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    texts.append(f"--- Page {i} ---\n{text}")
        return "\n\n".join(texts)
    except ImportError:
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            texts = []
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    texts.append(f"--- Page {i} ---\n{text}")
            return "\n\n".join(texts)
        except ImportError:
            raise ImportError("Install pdfplumber or PyPDF2: pip install pdfplumber")


def _extract_docx(file_path: str) -> str:
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs)
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")


def _extract_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_eml(file_path: str) -> str:
    import email
    from email import policy as email_policy

    with open(file_path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=email_policy.default)

    parts = []
    parts.append(f"From: {msg.get('From', 'N/A')}")
    parts.append(f"To: {msg.get('To', 'N/A')}")
    parts.append(f"CC: {msg.get('CC', 'N/A')}")
    parts.append(f"Date: {msg.get('Date', 'N/A')}")
    parts.append(f"Subject: {msg.get('Subject', 'N/A')}")
    parts.append("---")

    body = msg.get_body(preferencelist=("plain", "html"))
    if body:
        content = body.get_content()
        parts.append(content)

    return "\n".join(parts)


def _extract_csv(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        return "(empty CSV)"

    # Return header + first 100 rows as pipe-delimited text
    lines = []
    for i, row in enumerate(rows[:101]):
        lines.append(" | ".join(row))
        if i == 0:
            lines.append("-" * 40)

    if len(rows) > 101:
        lines.append(f"... ({len(rows) - 101} more rows)")

    return "\n".join(lines)


def _extract_xlsx_as_text(file_path: str) -> str:
    """Extract Excel file as text summary for non-data agents."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                parts.append(f"## Sheet: {sheet_name}\n(empty)")
                continue

            lines = [f"## Sheet: {sheet_name}"]
            for i, row in enumerate(rows[:101]):
                cells = [str(c) if c is not None else "" for c in row]
                lines.append(" | ".join(cells))
                if i == 0:
                    lines.append("-" * 40)

            if len(rows) > 101:
                lines.append(f"... ({len(rows) - 101} more rows)")
            parts.append("\n".join(lines))

        wb.close()
        return "\n\n".join(parts)
    except ImportError:
        raise ImportError("Install openpyxl: pip install openpyxl")


def get_excel_profile(file_path: str) -> dict:
    """Structured Excel profiling for the Excel agent. Returns dict with schema and stats."""
    try:
        import pandas as pd
    except ImportError:
        raise ImportError("Install pandas: pip install pandas")

    ext = Path(file_path).suffix.lower()
    if ext == ".csv":
        xls = {"Sheet1": pd.read_csv(file_path)}
    else:
        xls = pd.read_excel(file_path, sheet_name=None)

    profile = {}
    for sheet_name, df in xls.items():
        sheet_info = {
            "shape": list(df.shape),
            "columns": [],
        }
        for col in df.columns:
            col_info = {
                "name": str(col),
                "dtype": str(df[col].dtype),
                "null_pct": round(df[col].isna().mean() * 100, 1),
                "n_unique": int(df[col].nunique()),
            }
            if df[col].dtype in ("int64", "float64"):
                col_info["min"] = float(df[col].min()) if not df[col].isna().all() else None
                col_info["max"] = float(df[col].max()) if not df[col].isna().all() else None
                col_info["mean"] = round(float(df[col].mean()), 2) if not df[col].isna().all() else None
            else:
                top_vals = df[col].value_counts().head(5).to_dict()
                col_info["top_values"] = {str(k): int(v) for k, v in top_vals.items()}

            sheet_info["columns"].append(col_info)

        profile[sheet_name] = sheet_info

    return profile
